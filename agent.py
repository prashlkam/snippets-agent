
import requests
from bs4 import BeautifulSoup
import PyPDF2
import os
import re
from docx import Document
from urllib.parse import urlparse
from youtube_transcript_api._api import YouTubeTranscriptApi
from youtube_transcript_api._errors import NoTranscriptFound, TranscriptsDisabled

# Global constants
DOWNLOADS_DIR = os.path.join(os.path.expanduser("~"), "Downloads")
TEMP_DIR = os.path.join(DOWNLOADS_DIR, "temp")
OUTPUT_DOC_NAME = "Generated_Snippets.docx"

import google.generativeai as genai

# --- AI Generation ---

def get_ai_summary(content):
    """Generates a 60-word summary using an LLM."""
    # This is a placeholder. In a real scenario, you would use an LLM API.
    # For this example, we'll simulate the summary by truncating the content.
    words = content.split()
    summary = ' '.join(words[:60]) + '...'
    return summary

def get_seo_keywords(title):
    """Generates SEO keywords from a title using an LLM."""
    # Placeholder for SEO keyword generation
    # In a real scenario, you would use an LLM API.
    # For this example, we'll generate some dummy keywords.
    keywords = [word.lower() for word in title.split()[:5]]
    return keywords

from urllib.parse import urlparse, parse_qs

def extract_video_id_from_url(url):
    """Extracts the YouTube video ID from a URL."""
    parsed_url = urlparse(url)
    if "youtube.com" in parsed_url.netloc:
        query_params = parse_qs(parsed_url.query)
        if "v" in query_params:
            return query_params["v"][0]
    elif "youtu.be" in parsed_url.netloc:
        return parsed_url.path[1:]
    return None

def rewrite_summary_with_seo(summary, keywords):
    """Rewrites a summary to include SEO keywords using an LLM."""
    # Placeholder for summary rewriting
    # In a real scenario, you would use an LLM API.
    # For this example, we'll just append the keywords to the summary.
    return summary + " (Keywords: " + ', '.join(keywords) + ")"


# --- Output Generation ---

def save_as_word_doc(data, output_filename):
    """Saves the processed data to a Microsoft Word document."""
    doc = Document()
    doc.add_heading('AI Generated Snippets', level=1)

    for item in data:
        if item["status"] == "Processed":
            doc.add_heading(f"URL: {item['source_url']}", level=2)
            
            doc.add_heading('Title', level=3)
            doc.add_paragraph(item["title"])
            
            doc.add_heading('SEO Keywords', level=3)
            doc.add_paragraph(', '.join(item["seo_keywords"]))
            
            doc.add_heading('Final Summary', level=3)
            doc.add_paragraph(item["final_summary"])
            
            doc.add_page_break()

    try:
        doc.save(output_filename)
        print(f"\nSuccessfully saved snippets to: {output_filename}")
    except Exception as e:
        print(f"\nError saving Word document: {e}")


# --- Core Agent Functions ---


def get_content_type(url):
    """Checks if a URL points to a PDF, a web page, or a YouTube video."""
    if extract_video_id_from_url(url):
        return 'youtube'
    if url.lower().endswith('.pdf'):
        return 'pdf'
    try:
        response = requests.head(url, allow_redirects=True, timeout=10)
        content_type = response.headers.get('Content-Type', '').lower()
        if 'application/pdf' in content_type:
            return 'pdf'
        return 'html'
    except requests.RequestException:
        try:
            response = requests.get(url, stream=True, timeout=10)
            content_type = response.headers.get('Content-Type', '').lower()
            if 'application/pdf' in content_type:
                return 'pdf'
            return 'html'
        except requests.RequestException:
            return 'error'

def extract_html_content(url):
    """Fetches and extracts the title and main content from a web page."""
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        title = soup.find('title').get_text(strip=True) if soup.find('title') else "No Title Found"
        
        # A simple strategy to find the main content
        main_content = soup.find('article') or soup.find('main') or soup.find('div', class_='content') or soup.find('div', id='content')
        if main_content:
            content_text = main_content.get_text('\n', strip=True)
        else:
            content_text = soup.get_text('\n', strip=True) # Fallback to all text
            
        return title, content_text, None
    except requests.RequestException as e:
        return None, None, str(e)

def extract_youtube_content(url):
    """Fetches the transcript and title from a YouTube video."""
    try:
        video_id = extract_video_id_from_url(url)
        if not video_id:
            return None, None, "Invalid YouTube URL"

        # Get transcript
        transcript_list = YouTubeTranscriptApi().list(video_id)
        transcript = transcript_list.find_transcript(['en'])
        transcript_text = " ".join([item.text for item in transcript.fetch()])

        # Scrape title from YouTube page
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        title = soup.find('meta', property='og:title')
        title = title['content'] if title else "No Title Found"

        return title, transcript_text, None
    except NoTranscriptFound:
        return None, None, "No transcript found for this video."
    except TranscriptsDisabled:
        return None, None, "Transcripts are disabled for this video."
    
    except requests.RequestException as e:
        return None, None, f"Failed to fetch YouTube page: {e}"
    except Exception as e:
        return None, None, f"An unexpected error occurred: {e}"



def download_pdf(url, folder):
    """Downloads a PDF from a URL into a specified folder."""
    try:
        response = requests.get(url, stream=True, timeout=20)
        response.raise_for_status()
        
        # Get a filename from the URL
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        if not filename.lower().endswith('.pdf'):
            filename += ".pdf"
            
        filepath = os.path.join(folder, filename)
        
        with open(filepath, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        return filepath, None
    except requests.RequestException as e:
        return None, str(e)

def extract_pdf_content(filepath):
    """Extracts title and text content from a local PDF file."""
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Try to get title from metadata
            metadata = reader.metadata
            title = metadata.title if metadata and metadata.title else "No Title Found"

            # If no title in metadata, use the filename as a fallback
            if not title or title == "No Title Found":
                title = os.path.basename(filepath).replace('_', ' ').replace('-', ' ').rsplit('.', 1)[0]

            content_text = ""
            for page in reader.pages:
                content_text += page.extract_text() + "\n"
                
            return title, content_text, None
    except Exception as e:
        return None, None, str(e)

def process_urls(urls):
    """
    Main function to process a list of URLs and generate snippets.
    """
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)

    processed_data = []
    pdf_files_to_process = []

    for url in urls:
        print(f"Processing URL: {url}")
        content_type = get_content_type(url)
        
        url_data = {
            "source_url": url,
            "status": "Pending",
            "content_type": content_type,
            "raw_content": "",
            "title": "",
            "initial_summary": "",
            "seo_keywords": [],
            "final_summary": "",
            "error_message": ""
        }

        if content_type == 'html':
            print("Type: Web Page")
            title, content, error = extract_html_content(url)
            if error:
                url_data["status"] = "Error"
                url_data["error_message"] = error
            else:
                url_data["title"] = title
                url_data["raw_content"] = content
                url_data["status"] = "Content Extracted"
        elif content_type == 'youtube':
            print("Type: YouTube Video")
            title, content, error = extract_youtube_content(url)
            if error:
                url_data["status"] = "Error"
                url_data["error_message"] = error
            else:
                url_data["title"] = title
                url_data["raw_content"] = content
                url_data["status"] = "Content Extracted"
        elif content_type == 'pdf':
            print("Type: PDF Document")
            filepath, error = download_pdf(url, TEMP_DIR)
            if error:
                url_data["status"] = "Error"
                url_data["error_message"] = f"Failed to download PDF: {error}"
            else:
                print(f"PDF downloaded to: {filepath}")
                # We will process the PDF content after the initial loop
                pdf_files_to_process.append((url_data, filepath))
                url_data["status"] = "Downloaded"

        else:
            print("Type: Unknown or Error")
            url_data["status"] = "Error"
            url_data["error_message"] = "Could not determine content type or URL is unreachable."

        processed_data.append(url_data)
        print("-" * 20)

    # Now, process the downloaded PDFs
    for url_data, filepath in pdf_files_to_process:
        print(f"Extracting content from PDF: {filepath}")
        title, content, error = extract_pdf_content(filepath)
        if error:
            url_data["status"] = "Error"
            url_data["error_message"] = f"Failed to extract content from PDF: {error}"
        else:
            url_data["title"] = title
            url_data["raw_content"] = content
            url_data["status"] = "Content Extracted"
        print("-" * 20)


    # AI Processing: Summarization, Keyword Generation, and Rewriting
    for item in processed_data:
        if item["status"] == "Content Extracted":
            print(f"Generating AI content for: {item['source_url']}")
            
            # 1. Initial Summary
            item["initial_summary"] = get_ai_summary(item["raw_content"])
            
            # Clear raw content to save memory
            item["raw_content"] = ""
            
            # 2. SEO Keywords
            item["seo_keywords"] = get_seo_keywords(item["title"])
            
            # 3. Final Summary
            item["final_summary"] = rewrite_summary_with_seo(item["initial_summary"], item["seo_keywords"])
            
            item["status"] = "Processed"
            print("AI content generated.")
        print("-" * 20)

    # --- Final Output Generation ---
    output_path = os.path.join(os.getcwd(), OUTPUT_DOC_NAME)
    save_as_word_doc(processed_data, output_path)

    print("\nFinal Processed Data:")
    import json
    print(json.dumps(processed_data, indent=2))


# --- Main Execution ---

if __name__ == "__main__":
    # Placeholder list of URLs for testing
    sample_urls = [
        "https://www.youtube.com/watch?v=dQw4w9WgXcQ",
        "https://blog.google/technology/ai/google-gemini-ai/",
        "https://www.deeplearning.ai/the-batch/a-new-study-finds-that-llms-can-be-persuaded-to-give-up-private-data/",
        "https://arxiv.org/pdf/2305.15334.pdf" # Example PDF URL
    ]
    process_urls(sample_urls)
