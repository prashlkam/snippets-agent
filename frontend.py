import sys
import os
import time
from PyQt6.QtCore import QThread, pyqtSignal
from PyQt6.QtWidgets (
    QApplication, QMainWindow, QWidget, QVBoxLayout, 
    QTextEdit, QPushButton, QProgressBar, QTextBrowser, QMessageBox
)

# --- Import Agent Logic ---
# The following functions are copied and adapted from agent.py
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
from urllib.parse import urlparse
from youtube_transcript_api import YouTubeTranscriptApi

import google.generativeai as genai

# --- AI Model Initialization ---

def initialize_ai():
    """Initializes the generative AI model with an API key."""
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        # This is a fallback for local development, not recommended for production
        try:
            from api_key import GEMINI_API_KEY
            api_key = GEMINI_API_KEY
        except (ImportError, AttributeError):
            return None, "GEMINI_API_KEY not found in environment variables or api_key.py"
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        return model, None
    except Exception as e:
        return None, f"Failed to initialize AI model: {e}"

# Global constants
DOWNLOADS_DIR = os.path.join(os.path.expanduser("~"), "Downloads")
TEMP_DIR = os.path.join(DOWNLOADS_DIR, "temp")
OUTPUT_DOC_NAME = "Generated_Snippets.docx"

# --- AI Generation (Placeholders) ---
def get_ai_summary(content):
    words = content.split()
    summary = ' '.join(words[:60]) + '...'
    return summary

def get_seo_keywords(model, title):
    """Generates SEO keywords from a title using the Gemini API."""
    if not model:
        return ['AI model not initialized']
    try:
        prompt = f"Give top 5 SEO Keywords for \"{title}\". Return as a comma-separated list."
        response = model.generate_content(prompt)
        # Basic parsing, assuming the model returns a comma-separated string
        keywords = [kw.strip() for kw in response.text.split(',')]
        return keywords
    except Exception as e:
        return [f"Error generating keywords: {e}"]

import re

def rewrite_summary_with_seo(model, summary, keywords):
    """Rewrites a summary to include SEO keywords using the Gemini API."""
    if not model:
        return summary + " (Keywords: " + ', '.join(keywords) + ")" # Fallback
    try:
        kw_string = ", ".join(keywords)
        prompt = f"Rewrite the following summary to naturally include these SEO keywords: '{kw_string}'.\n\nSummary: '{summary}'\n\nDo not include any markdown formatting in your response."
        response = model.generate_content(prompt)
        # Remove common markdown characters
        clean_text = re.sub(r'[*_`#\[\]()]+', '', response.text)
        return clean_text.strip()
    except Exception as e:
        return summary + f" (Error rewriting summary: {e})"

# --- Agent Worker Thread ---
class AgentWorker(QThread):
    """Runs the snippet generation in a separate thread."""
    progress_update = pyqtSignal(int)
    log_update = pyqtSignal(str)
    finished = pyqtSignal(str) # Emits the path of the output file

    def __init__(self, urls):
        super().__init__()
        self.urls = urls
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        }

    def run(self):
        # --- Initialize AI Model ---
        model, error = initialize_ai()
        if error:
            self.log_update.emit(f"AI Initialization Error: {error}")
            # Decide if you want to stop or continue without AI features
            # For now, we'll log the error and let it continue, 
            # functions using the model will handle the `None` case.

        if not os.path.exists(TEMP_DIR):
            os.makedirs(TEMP_DIR)

        processed_data = []
        pdf_files_to_process = []
        total_urls = len(self.urls)

        for i, url in enumerate(self.urls):
            self.log_update.emit(f"({i+1}/{total_urls}) Processing URL: {url}")
            content_type = self.get_content_type(url)
            
            url_data = {
                "source_url": url, "status": "Pending", "content_type": content_type,
                "raw_content": "", "title": "", "initial_summary": "",
                "seo_keywords": [], "final_summary": "", "error_message": ""
            }

            if content_type == 'html':
                self.log_update.emit("-> Type: Web Page")
                title, content, error = self.extract_html_content(url)
                if error:
                    url_data["status"] = "Error"
                    url_data["error_message"] = error
                    self.log_update.emit(f"-> Error: {error}")
                else:
                    url_data["title"] = title
                    url_data["raw_content"] = content
                    url_data["status"] = "Content Extracted"
                    self.log_update.emit("-> Content extracted successfully.")
            elif content_type == 'pdf':
                self.log_update.emit("-> Type: PDF Document")
                filepath, error = self.download_pdf(url, TEMP_DIR)
                if error:
                    url_data["status"] = "Error"
                    url_data["error_message"] = f"Failed to download PDF: {error}"
                    self.log_update.emit(f"-> Error downloading PDF: {error}")
                else:
                    self.log_update.emit(f"-> PDF downloaded to: {filepath}")
                    pdf_files_to_process.append((url_data, filepath))
                    url_data["status"] = "Downloaded"
            elif content_type == 'youtube':
                self.log_update.emit("-> Type: YouTube Video")
                title, content, error = self.extract_youtube_content(url)
                if error:
                    url_data["status"] = "Error"
                    url_data["error_message"] = error
                    self.log_update.emit(f"-> Error: {error}")
                else:
                    url_data["title"] = title
                    url_data["raw_content"] = content
                    url_data["status"] = "Content Extracted"
                    self.log_update.emit("-> Content extracted successfully.")
            else:
                self.log_update.emit("-> Type: Unknown or Error")
                url_data["status"] = "Error"
                url_data["error_message"] = "Could not determine content type or URL is unreachable."

            processed_data.append(url_data)
            self.progress_update.emit(int(((i + 1) / total_urls) * 50)) # First 50% for download/extract

        # Process PDFs
        for url_data, filepath in pdf_files_to_process:
            self.log_update.emit(f"Extracting content from PDF: {os.path.basename(filepath)}")
            title, content, error = self.extract_pdf_content(filepath)
            if error:
                url_data["status"] = "Error"
                url_data["error_message"] = f"Failed to extract content from PDF: {error}"
                self.log_update.emit(f"-> Error extracting PDF content: {error}")
            else:
                url_data["title"] = title
                url_data["raw_content"] = content
                url_data["status"] = "Content Extracted"
                self.log_update.emit("-> Content extracted successfully.")

        # AI Processing
        total_items = len(processed_data)
        for i, item in enumerate(processed_data):
            if item["status"] == "Content Extracted":
                self.log_update.emit(f"Generating AI content for: {item['source_url']}")
                item["initial_summary"] = get_ai_summary(item["raw_content"])
                item["raw_content"] = ""
                item["seo_keywords"] = get_seo_keywords(model, item["title"])
                item["final_summary"] = rewrite_summary_with_seo(model, item["initial_summary"], item["seo_keywords"])
                item["status"] = "Processed"
                self.log_update.emit("-> AI content generated.")
            self.progress_update.emit(50 + int(((i + 1) / total_items) * 50)) # Second 50% for AI tasks

        # Save to Word Doc
        output_path = os.path.join(os.getcwd(), OUTPUT_DOC_NAME)
        self.save_as_word_doc(processed_data, output_path)
        self.finished.emit(output_path)

    def get_content_type(self, url):
        if 'youtube.com' in url or 'youtu.be' in url:
            return 'youtube'
        if url.lower().endswith('.pdf'): return 'pdf'
        try:
            res = requests.head(url, headers=self.headers, allow_redirects=True, timeout=10)
            ct = res.headers.get('Content-Type', '').lower()
            if 'application/pdf' in ct: return 'pdf'
        except requests.RequestException: return 'error'
        return 'html'

    def extract_html_content(self, url):
        try:
            res = requests.get(url, headers=self.headers, timeout=15)
            res.raise_for_status()
            soup = BeautifulSoup(res.content, 'html.parser')
            title = soup.find('title').get_text(strip=True) if soup.find('title') else "No Title"
            main = soup.find('article') or soup.find('main') or soup.find('div', class_='content')
            content = main.get_text('\n', strip=True) if main else soup.get_text('\n', strip=True)
            return title, content, None
        except requests.RequestException as e: return None, None, str(e)

    def extract_youtube_content(self, url):
        try:
            video_id = None
            if "youtube.com" in url:
                video_id = url.split("v=")[1].split("&")[0]
            elif "youtu.be" in url:
                video_id = url.split("/")[-1]

            if not video_id:
                return None, None, "Could not extract video ID from URL."

            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            transcript = " ".join([item["text"] for item in transcript_list])
            
            # Since we don't have a title from the transcript, 
            # we can try to fetch it from the youtube page, or just use the url
            return f"YouTube Video: {video_id}", transcript, None
        except Exception as e:
            return None, None, str(e)

    def download_pdf(self, url, folder):
        try:
            res = requests.get(url, headers=self.headers, stream=True, timeout=20)
            res.raise_for_status()
            fname = os.path.basename(urlparse(url).path)
            fpath = os.path.join(folder, fname)
            with open(fpath, 'wb') as f:
                for chunk in res.iter_content(chunk_size=8192): f.write(chunk)
            return fpath, None
        except requests.RequestException as e: return None, str(e)

    def extract_pdf_content(self, fpath):
        try:
            with open(fpath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                meta = reader.metadata
                title = meta.title if meta and meta.title else os.path.basename(fpath).rsplit('.', 1)[0]
                content = "".join(page.extract_text() for page in reader.pages)
                return title, content, None
        except Exception as e: return None, None, str(e)

    def save_as_word_doc(self, data, fname):
        doc = Document()
        doc.add_heading('AI Generated Snippets', level=1)
        for item in data:
            if item["status"] == "Processed":
                doc.add_heading(f"URL: {item['source_url']}", level=2)
                doc.add_heading('Title', level=3)
                doc.add_paragraph(item["title"])
                doc.add_heading('SEO Keywords', level=3)
                doc.add_paragraph(', '.join(item["seo_keywords"]))
                doc.add_heading('Final Summary', level=3)
                doc.add_paragraph(item["final_summary"])
                doc.add_page_break()
        try:
            doc.save(fname)
        except Exception as e:
            self.log_update.emit(f"Error saving Word doc: {e}")

# --- Frontend Class ---
class SnippetAgentFrontend(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Snippet Agent")
        self.setGeometry(100, 100, 800, 600)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        self.url_input = QTextEdit()
        self.url_input.setPlaceholderText("Enter a list of URLs, one per line...")
        layout.addWidget(self.url_input)

        self.start_button = QPushButton("Start Processing")
        layout.addWidget(self.start_button)

        self.progress_bar = QProgressBar()
        layout.addWidget(self.progress_bar)

        self.log_display = QTextBrowser()
        layout.addWidget(self.log_display)

        self.start_button.clicked.connect(self.start_processing)

    def start_processing(self):
        urls = self.url_input.toPlainText().strip().split('\n')
        if not urls or not urls[0]:
            QMessageBox.warning(self, "Input Error", "Please enter at least one URL.")
            return

        self.start_button.setEnabled(False)
        self.progress_bar.setValue(0)
        self.log_display.clear()
        self.log_display.append("Starting agent...")

        self.worker = AgentWorker(urls)
        self.worker.progress_update.connect(self.progress_bar.setValue)
        self.worker.log_update.connect(self.log_display.append)
        self.worker.finished.connect(self.on_processing_finished)
        self.worker.start()

    def on_processing_finished(self, output_path):
        self.log_display.append(f"\nProcessing complete! Output saved to: {output_path}")
        self.start_button.setEnabled(True)
        QMessageBox.information(self, "Success", f"Snippets saved to {output_path}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    frontend = SnippetAgentFrontend()
    frontend.show()
    sys.exit(app.exec())